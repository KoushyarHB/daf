# -*- coding: utf-8 -*-
"""python-docx helpers for DaF vocab cards.

``vocab.manifest.json`` holds canonical deck content. ``export_manifest_file`` (``manual-edit-from-docx`` / legacy ``pull``, ``export``)
merges a Word snapshot into the manifest.
``build_vocab_from_manifest_file`` (``manual-edit-from-manifest`` / legacy ``sync``,
``build``) regenerates Word from the manifest and rewrites normalized JSON.
"""

from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterator

from .audio_cards import assign_ids_to_deck, ensure_card_id, normalize_card_id
from .plural_forms import compact_plural_rule, format_plural_line, normalize_plural_fields
from .pos import normalize_pos
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DEFAULT_VOCAB_PATH = Path(__file__).resolve().parent.parent / "vocab.docx"
MANIFEST_PATH = Path(__file__).resolve().parent.parent / "vocab.manifest.json"

TITLE = "daf — vocabulary"
HEADING_BLUE = RGBColor(0x2F, 0x6F, 0xB8)
# Plural line: textbook ``die Antwort, -en`` (roman, same size as gloss).
IND = Inches(0.4)
NOTE_GRAY = RGBColor(0x40, 0x40, 0x40)
EXAMPLE_COLOR = RGBColor(0x30, 0x60, 0x8C)
# English gloss in parentheses after German in › lines (distinct from blue German):
EXAMPLE_TRANSLATION_COLOR = NOTE_GRAY

# Strong adjective suffix (after ``neu``) in optional grammar tables.
ADJ_SUFFIX_BLUE = RGBColor(0x1F, 0x4F, 0x8C)
GRAMMAR_HEADER_FILL = "EFF6FC"
GRAMMAR_CASE_COL_TWIPS = 680

_RE_NEU_ADJ = re.compile(r"\b(neu)(en|em|es|er|e)\b", re.IGNORECASE)

EXAMPLE_PREFIX = "\u203a "

DEFAULT_LEVEL = "A1"

# Legacy manifests used full ``die …`` NPs; map (lemma, plural NP) to textbook suffix fragment.
_LEGACY_LEMMA_PLURAL_TO_SUFFIX: dict[tuple[str, str], str] = {
    ("die W-Frage", "die W-Fragen"): "-n",
    ("die Antwort", "die Antworten"): "-en",
    ("die Tabelle", "die Tabellen"): "-n",
    ("die Verbform", "die Verbformen"): "-en",
    ("die Grammatik", "die Grammatiken"): "-en",
    ("der Blick", "die Blicke"): "-e",
    ("das Verb", "die Verben"): "-en",
    ("die Sekretärin", "die Sekretärinnen"): "-nen",
    ("Frau", "die Frauen"): "-en",
    ("die Frau", "die Frauen"): "-en",
    ("die Person", "die Personen"): "-en",
    ("das Foto", "die Fotos"): "-s",
    ("das Gespräch", "die Gespräche"): "-e",
    ("der Praktikant", "die Praktikanten"): "-en / -nen",
}


def parse_plural_line_from_word(text: str, head_line: str | None) -> tuple[str | None, str | None]:
    """Parse ``pluralRule · plural`` or legacy suffix-only plural lines from Word."""

    s = (text or "").strip()
    if not s or not head_line:
        return None, None

    if " · " in s:
        rule_part, form_part = [x.strip() for x in s.split(" · ", 1)]
        return compact_plural_rule(rule_part) or rule_part or None, form_part or None

    suffix = parse_plural_field_from_word_line(text, head_line)
    if suffix:
        return compact_plural_rule(suffix) or suffix, None
    return None, None


def parse_plural_field_from_word_line(text: str, head_line: str | None) -> str | None:
    """If an indented line is plural metadata, return the stored ``plural`` field (suffix side).

    Accepts textbook ``{lemma}, -en`` round-trip, shorthand ``→ -en``, or legacy ``lemma → die Plural`` from older Word.
    """

    s = (text or "").strip()
    if not s or not head_line:
        return None
    lemma_guess, _ipa = split_head(head_line.strip())
    lemma_key = lemma_guess.strip()

    arrow = "\u2192"
    if s.startswith(arrow):
        rhs = s[len(arrow) :].strip()
        return rhs or None
    if s.startswith("->"):
        rhs = s[2:].lstrip()
        return rhs or None

    for legacy_sep in (" \u2192 ", " → "):
        if legacy_sep in s:
            left, right = [x.strip() for x in s.split(legacy_sep, 1)]
            if left == lemma_key and right:
                return right

    if not lemma_key:
        return None
    for pref in (lemma_key + ", ", lemma_key + " , "):
        if s.startswith(pref):
            frag = s[len(pref) :].strip()
            return frag or None
    return None


def canonicalize_plural_field(head: str, raw: str) -> str:
    """Normalize ``plural`` JSON to textbook suffix notation (possibly with `` / …`` pairs)."""

    s = str(raw).strip()
    if not s:
        return ""
    lemma, _ = split_head(head.strip())
    lemma_key = lemma.strip()

    if s.startswith("-"):
        return s

    for legacy_sep in (" \u2192 ", " → "):
        if legacy_sep in s:
            left, right = [x.strip() for x in s.split(legacy_sep, 1)]
            if left == lemma_key and right:
                return canonicalize_plural_field(head, right)

    if s.startswith("die "):
        key = (lemma_key, s)
        if key in _LEGACY_LEMMA_PLURAL_TO_SUFFIX:
            return _LEGACY_LEMMA_PLURAL_TO_SUFFIX[key]
    return s


def utc_now_iso() -> str:
    """UTC timestamp for manifest metadata (ISO 8601, Z suffix)."""

    return datetime.now(timezone.utc).replace(microsecond=0).strftime("%Y-%m-%dT%H:%M:%SZ")


def _coerce_lektion(value: Any) -> int | None:
    if value is None or value is False:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, float) and value == int(value):
        return int(value)
    if isinstance(value, str) and not str(value).strip():
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def ordered_example_object(ex: dict[str, Any]) -> dict[str, Any]:
    """Stable keys: ``german``, ``english`` (null when no gloss), optional ``audio``."""

    g = (ex.get("german") or "").strip()
    e = ex.get("english")
    if e is None:
        en_out: str | None = None
    elif isinstance(e, str):
        en_out = e.strip() or None
    else:
        en_out = str(e).strip() or None
    out: dict[str, Any] = {"german": g, "english": en_out}
    audio = normalize_audio_field(ex.get("audio"))
    if audio:
        out["audio"] = audio
    return out


def normalize_grammar_table(raw: Any) -> dict[str, Any] | None:
    """Optional ``grammarTable`` block: ``{ \"columns\": [...], \"rows\": [[...], ...] }`` for Word + HTML."""

    if not isinstance(raw, dict):
        return None
    cols_raw = raw.get("columns")
    rows_raw = raw.get("rows")
    if not isinstance(cols_raw, list) or not cols_raw:
        return None
    if not isinstance(rows_raw, list) or not rows_raw:
        return None
    columns: list[str] = []
    for c in cols_raw:
        columns.append("" if c is None else str(c).strip())
    n = len(columns)
    if not n:
        return None
    rows: list[list[str]] = []
    for row in rows_raw:
        if not isinstance(row, list):
            return None
        cells: list[str] = []
        for i in range(n):
            v = row[i] if i < len(row) else None
            cells.append("" if v is None else str(v).strip())
        rows.append(cells)
    return {"columns": columns, "rows": rows}


def normalize_image_field(raw: Any) -> str | None:
    """Optional per-card website image path/URL."""

    if raw is None:
        return None
    s = str(raw).strip()
    return s or None


def normalize_audio_field(raw: Any) -> str | None:
    """Optional per-card pronunciation audio path (e.g. ``/audio/lemma.mp3``)."""

    if raw is None:
        return None
    s = str(raw).strip()
    return s or None


def iter_grammar_adj_suffix_runs(text: str) -> Iterator[tuple[str, bool]]:
    """Split cell text into runs; booleans mark the adjective ending right after ``neu``."""

    if not text:
        return
    yield from _iter_grammar_adj_suffix_runs_impl(text)


def _iter_grammar_adj_suffix_runs_impl(text: str) -> Iterator[tuple[str, bool]]:
    last = 0
    matched = False
    for m in _RE_NEU_ADJ.finditer(text):
        matched = True
        if m.start() > last:
            yield text[last:m.start()], False
        yield m.group(1), False
        yield m.group(2), True
        last = m.end()
    if not matched:
        yield text, False
    elif last < len(text):
        yield text[last:], False


def _set_tbl_left_indent(tbl, twips: int) -> None:
    tbl_pr = tbl._tbl.tblPr
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(twips))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)


def _set_cell_width_twips(tc_pr, twips: int) -> None:
    tag = qn("w:tcW")
    for child in list(tc_pr):
        if child.tag == tag:
            tc_pr.remove(child)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _set_cell_fill_shading(tc_pr, fill: str) -> None:
    tag_shd = qn("w:shd")
    for child in list(tc_pr):
        if child.tag == tag_shd:
            tc_pr.remove(child)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _fill_docx_grammar_cell(
    cell,
    text: str,
    *,
    point_size: float,
    role: str,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)

    if role == "hdr":
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(point_size)
        r.bold = True
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        _set_cell_fill_shading(tc_pr, GRAMMAR_HEADER_FILL)
        return

    if role == "case":
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(point_size)
        r.font.color.rgb = NOTE_GRAY
        return

    for chunk, is_suffix in iter_grammar_adj_suffix_runs(text):
        if chunk == "":
            continue
        run = p.add_run(chunk)
        run.font.name = "Calibri"
        run.font.size = Pt(point_size)
        run.font.color.rgb = ADJ_SUFFIX_BLUE if is_suffix else NOTE_GRAY
        if is_suffix:
            run.bold = True


def ordered_manifest_card(c: dict[str, Any]) -> dict[str, Any]:
    """Stable key order for JSON: content fields, then metadata."""

    ex_raw = c.get("examples") or []
    examples_ord = [ordered_example_object(x) for x in ex_raw if isinstance(x, dict)]
    head, ipa = normalize_head_ipa_fields(c)
    plural_rule, plural_form = normalize_plural_fields(c, head)
    od: dict[str, Any] = {"head": head}
    if ipa:
        od["ipa"] = ipa
    cid = normalize_card_id(c.get("id"))
    if cid:
        od["id"] = cid
    od["pos"] = normalize_pos(c)
    if plural_rule:
        od["pluralRule"] = plural_rule
    if plural_form:
        od["plural"] = plural_form
    od["gloss"] = c["gloss"]
    od["notes"] = c["notes"]
    grammar_table = normalize_grammar_table(c.get("grammarTable"))
    if grammar_table:
        od["grammarTable"] = grammar_table
    image = normalize_image_field(c.get("image"))
    if image:
        od["image"] = image
    audio = normalize_audio_field(c.get("audio"))
    if audio:
        od["audio"] = audio
    od["examples"] = examples_ord
    od["createdAt"] = c["createdAt"]
    od["updatedAt"] = c["updatedAt"]
    od["lektion"] = c["lektion"]
    od["level"] = c["level"]
    od["studied"] = bool(c.get("studied"))
    return od


def english_gloss_inner(en: str | None) -> str | None:
    """Canonical manifest stores English without wrapping parentheses; migrate legacy ``(English)``."""

    if en is None:
        return None
    s = str(en).strip()
    if not s:
        return None
    if len(s) >= 2 and s[0] == "(" and s[-1] == ")":
        return s[1:-1].strip() or None
    return s


def _examples_from_legacy_string(s: str) -> list[dict[str, Any]]:
    """One legacy manifest string (possibly with `` / `` between units) → example objects."""

    s = s.strip()
    if not s:
        return []
    out: list[dict[str, Any]] = []
    for chunk in split_example_body_units(s):
        de, en = split_example_chunk_translation(chunk)
        de = (de or "").strip()
        en_in = english_gloss_inner(en) if en else None
        if de or en_in:
            out.append({"german": de, "english": en_in})
    return out


def normalize_examples_from_card(card: dict[str, Any]) -> list[dict[str, Any]]:
    """Canonical ``examples``: ``[{ \"german\", \"english\" }, ...]``; migrate strings and ``example_units``."""

    out: list[dict[str, Any]] = []
    units = card.get("example_units")
    if isinstance(units, list) and units:
        for pair in units:
            if not isinstance(pair, (list, tuple)) or len(pair) < 2:
                continue
            g = str(pair[0]).strip()
            raw_e = str(pair[1]).strip() if pair[1] is not None else ""
            e = english_gloss_inner(raw_e) if raw_e else None
            if g or e:
                out.append({"german": g, "english": e})
    for item in card.get("examples") or []:
        if isinstance(item, dict):
            g = (item.get("german") or "").strip()
            e = item.get("english")
            if e is None:
                en_out = None
            elif isinstance(e, str):
                en_out = english_gloss_inner(e)
            else:
                en_out = english_gloss_inner(str(e))
            if g or en_out:
                blob: dict[str, Any] = {"german": g, "english": en_out}
                aud = normalize_audio_field(item.get("audio"))
                if aud:
                    blob["audio"] = aud
                out.append(ordered_example_object(blob))
        elif isinstance(item, str):
            out.extend(_examples_from_legacy_string(item))
    return out


def normalize_card_meta(
    card: dict[str, Any],
    *,
    used_ids: set[str] | None = None,
) -> dict[str, Any]:
    """Fill required manifest metadata; coerce types. Word export code paths should call this before writing JSON."""

    now = utc_now_iso()
    head, ipa = normalize_head_ipa_fields(card)
    if used_ids is not None:
        ensure_card_id(card, used_ids)
    gloss = [x for x in (card.get("gloss") or []) if isinstance(x, str)]
    notes = [x for x in (card.get("notes") or []) if isinstance(x, str)]
    examples = normalize_examples_from_card(card)
    created = card.get("createdAt")
    updated = card.get("updatedAt")
    if not created:
        created = updated if updated else now
    if not updated:
        updated = created if created else now
    level_raw = card.get("level")
    level = str(level_raw).strip() if level_raw else DEFAULT_LEVEL
    if not level:
        level = DEFAULT_LEVEL
    plural_raw = card.get("plural")
    plural_rule_raw = card.get("pluralRule")
    plural_rule, plural_form = normalize_plural_fields(
        {
            **card,
            "head": head,
            "pluralRule": plural_rule_raw,
            "plural": plural_raw,
        },
        head,
    )
    cid = normalize_card_id(card.get("id"))
    pos = normalize_pos({**card, "head": head})
    out: dict[str, Any] = {
        "head": head,
        "pos": pos,
        "gloss": gloss,
        "notes": notes,
        "examples": examples,
        "createdAt": str(created),
        "updatedAt": str(updated),
        "lektion": _coerce_lektion(card.get("lektion")),
        "level": level,
        "studied": bool(card.get("studied")),
    }
    if ipa:
        out["ipa"] = ipa
    if cid:
        out["id"] = cid
    if plural_rule:
        out["pluralRule"] = plural_rule
    if plural_form:
        out["plural"] = plural_form
    grammar_table = normalize_grammar_table(card.get("grammarTable"))
    if grammar_table:
        out["grammarTable"] = grammar_table
    image = normalize_image_field(card.get("image"))
    if image:
        out["image"] = image
    audio = normalize_audio_field(card.get("audio"))
    if audio:
        out["audio"] = audio
    return ordered_manifest_card(out)


def merge_parsed_cards_with_previous_manifest(
    parsed: list[dict[str, Any]],
    previous: list[dict[str, Any]] | None,
) -> list[dict[str, Any]]:
    """Attach metadata from an existing manifest (matched by ``head``). New heads get fresh timestamps and defaults."""

    now = utc_now_iso()
    prev_by_head = _index_previous_cards(previous)
    out: list[dict[str, Any]] = []
    used_ids: set[str] = set()
    for c in parsed:
        h = (c.get("head") or "").strip()
        ipa_parsed = str(c.get("ipa") or "").strip() or None
        prev = None
        for key in _head_lookup_keys(h, ipa_parsed):
            prev = prev_by_head.get(key)
            if prev:
                break
        merged: dict[str, Any] = {
            "head": h,
            "gloss": list(c.get("gloss") or []),
            "notes": list(c.get("notes") or []),
            "examples": normalize_examples_from_card(c),
        }
        if ipa_parsed:
            merged["ipa"] = ipa_parsed
        gt_new = normalize_grammar_table(c.get("grammarTable"))
        if gt_new:
            merged["grammarTable"] = gt_new
        elif prev:
            gt_keep = normalize_grammar_table(prev.get("grammarTable"))
            if gt_keep:
                merged["grammarTable"] = gt_keep
        image_new = normalize_image_field(c.get("image"))
        if image_new:
            merged["image"] = image_new
        elif prev:
            image_keep = normalize_image_field(prev.get("image"))
            if image_keep:
                merged["image"] = image_keep
        audio_new = normalize_audio_field(c.get("audio"))
        if audio_new:
            merged["audio"] = audio_new
        elif prev:
            audio_keep = normalize_audio_field(prev.get("audio"))
            if audio_keep:
                merged["audio"] = audio_keep
        id_new = normalize_card_id(c.get("id"))
        if id_new:
            merged["id"] = id_new
        elif prev:
            id_keep = normalize_card_id(prev.get("id"))
            if id_keep:
                merged["id"] = id_keep
        pl_new = (c.get("plural") or "").strip()
        if pl_new:
            merged["plural"] = pl_new
        elif prev:
            pl_keep = (prev.get("plural") or "").strip()
            if pl_keep:
                merged["plural"] = pl_keep
        pl_rule_new = (c.get("pluralRule") or "").strip()
        if pl_rule_new:
            merged["pluralRule"] = pl_rule_new
        elif prev:
            pl_rule_keep = (prev.get("pluralRule") or "").strip()
            if pl_rule_keep:
                merged["pluralRule"] = pl_rule_keep
        pos_new = str(c.get("pos") or "").strip().lower()
        if pos_new:
            merged["pos"] = pos_new
        elif prev:
            pos_keep = str(prev.get("pos") or "").strip().lower()
            if pos_keep:
                merged["pos"] = pos_keep
        if prev:
            merged["createdAt"] = prev.get("createdAt") or now
            merged["updatedAt"] = now
            merged["lektion"] = _coerce_lektion(prev.get("lektion"))
            merged["level"] = str(prev.get("level") or DEFAULT_LEVEL).strip() or DEFAULT_LEVEL
            merged["studied"] = bool(prev.get("studied"))
        else:
            merged["createdAt"] = now
            merged["updatedAt"] = now
            merged["lektion"] = None
            merged["level"] = DEFAULT_LEVEL
            merged["studied"] = False
        out.append(normalize_card_meta(merged, used_ids=used_ids))
    return out


def _parse_head_card_start(p) -> dict[str, Any]:
    """Parse a card head paragraph into ``head`` + optional ``ipa`` (bold vs roman runs)."""

    bold_parts: list[str] = []
    ipa_parts: list[str] = []
    for run in p.runs:
        text = run.text or ""
        if not text:
            continue
        if run.bold:
            bold_parts.append(text)
        else:
            ipa_parts.append(text)
    head = "".join(bold_parts).strip()
    ipa_raw = "".join(ipa_parts).strip()
    if head and ipa_raw:
        ipa = normalize_ipa_storage(ipa_raw)
        card: dict[str, Any] = {"head": head, "gloss": [], "notes": [], "examples": []}
        if ipa:
            card["ipa"] = ipa
        return card
    full = (p.text or "").strip()
    lemma, ipa = split_head_and_ipa(full)
    card: dict[str, Any] = {"head": lemma, "gloss": [], "notes": [], "examples": []}
    if ipa:
        card["ipa"] = ipa
    return card


def indented(p) -> bool:
    return bool(p.paragraph_format.left_indent)


def _is_example_paragraph(text: str) -> bool:
    """Examples are indented lines that begin with ``›`` (U+203A); glosses never use this marker."""

    s = text.lstrip()
    if not s:
        return False
    if s.startswith("\u203a"):
        return True
    return s.startswith("›")


def role(p) -> str:
    """Classify paragraphs for vocabulary cards (titles + separators ignored upstream)."""

    if p.style and p.style.name == "Heading 1":
        return "title"
    t = p.text or ""
    if not t.strip():
        return "blank"
    if t.strip() == TITLE:
        return "title"
    if not indented(p):
        return "head"
    tl = t.lstrip()
    if tl.startswith("\u2022") or tl.startswith("•"):
        return "note"
    if _is_example_paragraph(t):
        return "example"
    return "gloss"


def parse_vocab_document(doc_path: Path) -> list[dict[str, Any]]:
    """Turn an existing vocab Word file into normalized dict cards for JSON export."""

    doc = Document(str(doc_path))
    cards: list[dict[str, Any]] = []
    cur: dict[str, Any] | None = None

    for p in doc.paragraphs:
        rr = role(p)
        if rr in {"blank", "title"}:
            continue
        if rr == "head":
            if cur:
                cards.append(cur)
            cur = _parse_head_card_start(p)
            continue
        if cur is None:
            continue
        txt = (p.text or "").strip()
        if rr == "gloss":
            pl_rule, pl_form = parse_plural_line_from_word(txt, cur.get("head"))
            if pl_rule or pl_form:
                if pl_rule:
                    cur["pluralRule"] = pl_rule
                if pl_form:
                    cur["plural"] = pl_form
                continue
            cur["gloss"].append(txt)
        elif rr == "note":
            cur["notes"].append(txt)
        else:
            body = strip_example_surface(txt)
            for chunk in split_example_body_units(body):
                de, en = split_example_chunk_translation(chunk)
                de = (de or "").strip()
                en_in = english_gloss_inner(en) if en else None
                if de or en_in:
                    cur.setdefault("examples", []).append({"german": de, "english": en_in})

    if cur:
        cards.append(cur)
    return cards


def strip_example_surface(text: str) -> str:
    """Keep only the textual body following the leading › marker."""

    s = text.lstrip()
    if s.startswith("\u203a"):
        return s[len("\u203a") :].lstrip()
    if s.startswith("›"):
        return s[1:].lstrip()
    return text


def format_example_chunks(pairs: list[tuple[str, str]]) -> str:
    """`(de, en)` → `DE (English) / DE2 (English2)` respecting the project's example rule."""

    parts: list[str] = []
    for de_raw, en_raw in pairs:
        chunk = f"{de_raw.strip()} ({en_raw.strip()})"
        parts.append(chunk)
    return " / ".join(parts)


_IPA_MARKERS = "ˈˌɪʊəɐ̯ːʃçɡɪɛɔʁ̩ʔ."
_RE_IPA_TOKEN = re.compile(
    rf" /(?=[^/]*[{re.escape(_IPA_MARKERS)}])[^/]+/"
)


def _ipa_tokens_to_storage(matches: list[str]) -> str | None:
    parts: list[str] = []
    for token in matches:
        inner = token.strip()[1:-1].strip("/").strip()
        if inner:
            parts.append(inner)
    return " ".join(parts) if parts else None


def normalize_ipa_storage(raw: str | None) -> str | None:
    """Canonical manifest ``ipa``: bare phonetic text without surrounding slashes."""

    if raw is None or not str(raw).strip():
        return None
    s = str(raw).strip()
    matches = _RE_IPA_TOKEN.findall(s)
    if matches:
        return _ipa_tokens_to_storage(matches)
    bare = s.strip("/").strip()
    return bare or None


def format_ipa_display(ipa: str | None) -> str | None:
    """Word/HTML suffix: `` /ɡə…/`` (or multiple slash-wrapped tokens)."""

    bare = normalize_ipa_storage(ipa)
    if not bare:
        return None
    return " " + " ".join(f"/{part}/" for part in bare.split())


def split_head_and_ipa(combined: str) -> tuple[str, str | None]:
    """Split lemma from trailing IPA token(s) like `` /ɡə…/`` (not gender `` / die …``)."""

    s = combined.strip()
    if not s:
        return "", None
    matches = list(_RE_IPA_TOKEN.finditer(s))
    if not matches:
        return s, None
    lemma = s[: matches[0].start()].rstrip()
    tail = s[matches[0].start() :]
    storage = _ipa_tokens_to_storage(_RE_IPA_TOKEN.findall(tail))
    return lemma or s, storage


def normalize_head_ipa_fields(card: dict[str, Any]) -> tuple[str, str | None]:
    """Return canonical ``(head, ipa)``; migrate legacy combined ``head`` strings."""

    raw_head = str(card.get("head") or "").strip()
    raw_ipa = card.get("ipa")
    ipa_explicit = normalize_ipa_storage(str(raw_ipa).strip() if raw_ipa else None)
    if ipa_explicit:
        lemma, _trailing = split_head_and_ipa(raw_head)
        return lemma, ipa_explicit
    return split_head_and_ipa(raw_head)


def split_head(head: str) -> tuple[str, str | None]:
    """Lemma and optional IPA from a head line (legacy combined or lemma-only)."""

    lemma, ipa = split_head_and_ipa(head)
    return lemma, ipa


def _head_lookup_keys(head: str, ipa: str | None = None) -> set[str]:
    keys: set[str] = set()
    h = head.strip()
    if h:
        keys.add(h)
    if ipa:
        keys.add(f"{h} {ipa}".strip())
        display = format_ipa_display(ipa)
        if display:
            keys.add(f"{h}{display}".strip())
    lemma, legacy_ipa = split_head_and_ipa(h)
    if legacy_ipa:
        keys.add(h)
        keys.add(lemma)
        keys.add(f"{lemma} {legacy_ipa}".strip())
        legacy_display = format_ipa_display(legacy_ipa)
        if legacy_display:
            keys.add(f"{lemma}{legacy_display}".strip())
    return keys


def _index_previous_cards(previous: list[dict[str, Any]] | None) -> dict[str, dict[str, Any]]:
    prev_by_head: dict[str, dict[str, Any]] = {}
    if not previous:
        return prev_by_head
    for c in previous:
        if not isinstance(c, dict) or not c.get("head"):
            continue
        h = str(c["head"]).strip()
        ipa = str(c.get("ipa") or "").strip() or None
        for key in _head_lookup_keys(h, ipa):
            prev_by_head.setdefault(key, c)
        lemma, legacy_ipa = split_head_and_ipa(h)
        if legacy_ipa and not ipa:
            combined = f"{lemma} {legacy_ipa}".strip()
            prev_by_head.setdefault(combined, c)
    return prev_by_head


def _new_base_document():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(TITLE, level=1)
    h.paragraph_format.space_after = Pt(12)
    for run in h.runs:
        run.font.color.rgb = HEADING_BLUE
    return doc


def split_example_chunk_translation(chunk: str) -> tuple[str, str | None]:
    """Split ``German (English)`` into German text and ``(English)``, or return whole chunk as German."""

    chunk = chunk.strip()
    if not chunk:
        return "", None
    if chunk.endswith(")") and " (" in chunk:
        i = chunk.rfind(" (")
        if i >= 0:
            return chunk[:i].strip(), chunk[i + 1 :].strip()
    return chunk, None


def split_example_body_units(body: str) -> list[str]:
    """Split merged example text on `` / `` only outside ``(...)``.

    English inside parentheses may contain `` / `` (e.g. *Good afternoon / hello*).
    """

    s = body.strip()
    if not s:
        return []
    units: list[str] = []
    start = 0
    depth = 0
    n = len(s)
    i = 0
    while i < n:
        c = s[i]
        if c == "(":
            depth += 1
        elif c == ")":
            depth = max(0, depth - 1)
        elif depth == 0 and i + 3 <= n and s[i : i + 3] == " / ":
            units.append(s[start:i].strip())
            start = i + 3
            i += 3
            continue
        i += 1
    units.append(s[start:].strip())
    return [u for u in units if u]


def add_example_objects(doc: Document, objects: list[dict[str, Any]]) -> None:
    """One indented ``›`` paragraph per example (German slate blue; optional ``(English)`` gray)."""

    objs: list[dict[str, Any]] = []
    for obj in objects:
        de = (obj.get("german") or "").strip()
        e = obj.get("english")
        if e is None:
            en: str | None = None
        elif isinstance(e, str):
            en = e.strip() or None
        else:
            en = str(e).strip() or None
        if de or en:
            objs.append({"german": de, "english": en})
    n = len(objs)
    if not n:
        return
    for idx, obj in enumerate(objs):
        de = obj["german"]
        en = obj["english"]
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = IND
        paragraph.paragraph_format.space_before = Pt(2) if idx == 0 else Pt(0)
        paragraph.paragraph_format.space_after = Pt(6) if idx == n - 1 else Pt(3)

        for piece, color in (
            (EXAMPLE_PREFIX, EXAMPLE_COLOR),
            (de, EXAMPLE_COLOR),
        ):
            run = paragraph.add_run(piece)
            run.font.size = Pt(10.5)
            run.font.italic = True
            run.font.color.rgb = color
        if en:
            run_en = paragraph.add_run(" (" + en + ")")
            run_en.font.size = Pt(10.5)
            run_en.font.italic = True
            run_en.font.color.rgb = EXAMPLE_TRANSLATION_COLOR


def add_grammar_table_word(doc: Document, grammar_table: dict[str, Any]) -> None:
    cols = grammar_table.get("columns") or []
    rows = grammar_table.get("rows") or []
    if not cols or not rows:
        return
    nc = len(cols)
    table = doc.add_table(rows=len(rows) + 1, cols=nc)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    _set_tbl_left_indent(table, int(IND.twips))
    narrow_first_col = bool(cols and str(cols[0]).strip() == "")

    hdr_cells = table.rows[0].cells
    for j in range(nc):
        tc_pr = hdr_cells[j]._tc.get_or_add_tcPr()
        if narrow_first_col and j == 0:
            _set_cell_width_twips(tc_pr, GRAMMAR_CASE_COL_TWIPS)
        _fill_docx_grammar_cell(
            hdr_cells[j],
            cols[j],
            point_size=10,
            role="hdr",
        )

    for ri, row_vals in enumerate(rows, start=1):
        row_cells = table.rows[ri].cells
        for j in range(nc):
            tc_pr = row_cells[j]._tc.get_or_add_tcPr()
            if narrow_first_col and j == 0:
                _set_cell_width_twips(tc_pr, GRAMMAR_CASE_COL_TWIPS)
            txt = row_vals[j] if j < len(row_vals) else ""
            role = "case" if j == 0 else "phrase"
            _fill_docx_grammar_cell(
                row_cells[j],
                txt,
                point_size=9.5,
                role=role,
            )
    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(6)


def write_card(doc: Document, card: dict[str, Any], *, is_first: bool):
    if not is_first:
        spacer = doc.add_paragraph("")
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    head, ipa = normalize_head_ipa_fields(card)
    head_p = doc.add_paragraph()
    head_p.paragraph_format.space_after = Pt(6)
    head_p.add_run(head).bold = True
    ipa_display = format_ipa_display(ipa)
    if ipa_display:
        r_ipa = head_p.add_run(ipa_display)
        r_ipa.bold = False
        r_ipa.font.size = Pt(10)
        r_ipa.font.color.rgb = NOTE_GRAY

    plural_rule, plural_form = normalize_plural_fields(card, head)
    if plural_rule and plural_form:
        plural_line = format_plural_line(plural_rule, plural_form)
        pl_p = doc.add_paragraph()
        pl_p.paragraph_format.left_indent = IND
        pl_p.paragraph_format.space_after = Pt(4)
        r_pl = pl_p.add_run(plural_line)
        r_pl.font.size = Pt(11)

    for g in card.get("gloss") or []:
        g_p = doc.add_paragraph(g)
        g_p.paragraph_format.left_indent = IND
        g_p.paragraph_format.space_after = Pt(4)

    for n in card.get("notes") or []:
        note_p = doc.add_paragraph()
        note_p.paragraph_format.left_indent = IND
        note_p.paragraph_format.space_after = Pt(4)
        r = note_p.add_run(n)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = NOTE_GRAY

    grammar_block = normalize_grammar_table(card.get("grammarTable"))
    if grammar_block:
        add_grammar_table_word(doc, grammar_block)

    add_example_objects(doc, normalize_examples_from_card(card))


def build_vocab_document(cards: list[dict[str, Any]]) -> Document:
    doc = _new_base_document()
    for idx, card in enumerate(cards):
        write_card(doc, card, is_first=idx == 0)
    return doc


def rebuild_vocab_layout(source_path: Path, destination_path: Path | None = None) -> Path:
    """Parse + rewrite to normalize spacing/example styling."""

    dst = destination_path or source_path
    cards = parse_vocab_document(source_path)
    doc = build_vocab_document(cards)
    doc.save(str(dst))
    return dst


def export_manifest_file(
    vocab_path: Path = DEFAULT_VOCAB_PATH,
    manifest_path: Path = MANIFEST_PATH,
) -> Path:
    """Merge a Word snapshot into ``manifest_path`` (canonical JSON on disk)."""

    parsed = parse_vocab_document(vocab_path)
    previous: list[dict[str, Any]] | None = None
    if manifest_path.exists():
        try:
            raw = json.loads(manifest_path.read_text(encoding="utf-8"))
            if isinstance(raw, list):
                previous = raw
        except (json.JSONDecodeError, OSError):
            previous = None
    data = merge_parsed_cards_with_previous_manifest(parsed, previous)
    manifest_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest_path


def build_vocab_from_manifest_file(
    manifest_path: Path = MANIFEST_PATH,
    vocab_path: Path = DEFAULT_VOCAB_PATH,
) -> Path:
    blob = json.loads(manifest_path.read_text(encoding="utf-8"))
    if not isinstance(blob, list):
        raise ValueError("Manifest must be a JSON array of card objects")
    used_ids: set[str] = set()
    cards = [
        normalize_card_meta(c, used_ids=used_ids) for c in blob if isinstance(c, dict)
    ]
    if len(cards) != len(blob):
        raise ValueError("Manifest must contain only JSON objects (card dicts)")
    doc = build_vocab_document(cards)
    doc.save(str(vocab_path))
    manifest_path.write_text(json.dumps(cards, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    from .html_preview import write_vocab_preview

    write_vocab_preview(manifest_path)
    return vocab_path


def set_card_studied_flag(manifest_path: Path, card_id: str, studied: bool) -> bool:
    """Update one card's ``studied`` flag in ``vocab.manifest.json`` (matched by ``id``)."""

    card_id = str(card_id or "").strip()
    if not card_id:
        return False
    manifest_path = Path(manifest_path)
    try:
        blob = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return False
    if not isinstance(blob, list):
        return False
    found = False
    for raw in blob:
        if not isinstance(raw, dict):
            continue
        cid = normalize_card_id(raw.get("id"))
        if cid != card_id:
            continue
        raw["studied"] = bool(studied)
        raw["updatedAt"] = utc_now_iso()
        found = True
        break
    if not found:
        return False
    used_ids: set[str] = set()
    cards = [normalize_card_meta(c, used_ids=used_ids) for c in blob if isinstance(c, dict)]
    manifest_path.write_text(json.dumps(cards, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return True
